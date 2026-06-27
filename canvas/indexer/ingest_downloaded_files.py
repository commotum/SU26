#!/usr/bin/env python3
"""Copy already downloaded Canvas files into the repo and parse text evidence."""

from __future__ import annotations

import argparse
import csv
import hashlib
import re
import shutil
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".rtf"}


@dataclass
class DownloadGroup:
    course_id: str
    course_code: str
    canvas_file_id: str
    file_names: set[str] = field(default_factory=set)
    download_ids: set[str] = field(default_factory=set)
    download_manifest_ids: set[str] = field(default_factory=set)
    object_ids: set[str] = field(default_factory=set)
    hrefs: set[str] = field(default_factory=set)
    source_page_ids: set[str] = field(default_factory=set)
    source_page_titles: set[str] = field(default_factory=set)


def clean_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def substantive_text(value: str) -> str:
    without_page_labels = re.sub(r"(?im)^## Page \d+\s*", "", value)
    return clean_text(without_page_labels)


def safe_file_name(value: str) -> str:
    name = Path(value).name
    stem = re.sub(r"[^A-Za-z0-9._ -]+", "-", name).strip(" .-")
    return stem or "downloaded-file"


def normalized_name(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", Path(value).name.lower())


def csv_join(values: Iterable[str]) -> str:
    return "|".join(sorted({clean_text(value) for value in values if clean_text(value)}))


def sha256_file(file_path: Path) -> str:
    digest = hashlib.sha256()
    with file_path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def destination_path_for(source_path: Path, downloaded_dir: Path, group: DownloadGroup) -> Path:
    prefix = f"{group.course_code}_{group.canvas_file_id}_"
    if source_path.parent.resolve() == downloaded_dir.resolve() and source_path.name.startswith(prefix):
        return source_path
    return downloaded_dir / f"{prefix}{safe_file_name(source_path.name)}"


def read_csv(file_path: Path) -> list[dict[str, str]]:
    with file_path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def read_text_overrides(file_path: Path) -> dict[tuple[str, str], dict[str, str]]:
    if not file_path.exists():
        return {}
    overrides: dict[tuple[str, str], dict[str, str]] = {}
    for row in read_csv(file_path):
        course_code = clean_text(row.get("course_code"))
        canvas_file_id = clean_text(row.get("canvas_file_id"))
        text = str(row.get("text") or "").strip()
        if course_code and canvas_file_id and text:
            overrides[(course_code, canvas_file_id)] = row
    return overrides


def write_csv(file_path: Path, rows: list[dict[str, object]], fieldnames: list[str]) -> None:
    file_path.parent.mkdir(parents=True, exist_ok=True)
    with file_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fieldnames})


def build_download_groups(download_manifest_rows: list[dict[str, str]]) -> dict[tuple[str, str], DownloadGroup]:
    groups: dict[tuple[str, str], DownloadGroup] = {}
    for row in download_manifest_rows:
        canvas_file_id = clean_text(row.get("canvas_file_id"))
        course_code = clean_text(row.get("course_code"))
        if not canvas_file_id or not course_code:
            continue
        key = (course_code, canvas_file_id)
        if key not in groups:
            groups[key] = DownloadGroup(
                course_id=clean_text(row.get("course_id")),
                course_code=course_code,
                canvas_file_id=canvas_file_id,
            )
        group = groups[key]
        group.file_names.add(clean_text(row.get("file_name")))
        group.download_ids.add(clean_text(row.get("download_id")))
        group.download_manifest_ids.add(clean_text(row.get("download_manifest_id")))
        group.object_ids.add(clean_text(row.get("object_id")))
        group.hrefs.add(clean_text(row.get("href")))
        group.source_page_ids.add(clean_text(row.get("source_page_id")))
        group.source_page_titles.add(clean_text(row.get("source_page_title")))
    return groups


def match_score(local_file: Path, group: DownloadGroup) -> int:
    local_name = local_file.name.lower()
    local_norm = normalized_name(local_file.name)
    if local_file.name.startswith(f"{group.course_code}_{group.canvas_file_id}_"):
        return 100
    best = 0
    for manifest_name in group.file_names:
        if not manifest_name or "." not in manifest_name:
            continue
        manifest_lower = manifest_name.lower()
        manifest_norm = normalized_name(manifest_name)
        if local_name == manifest_lower:
            best = max(best, 100)
        elif local_name.endswith(manifest_lower):
            best = max(best, 90)
        elif local_norm == manifest_norm:
            best = max(best, 85)
        elif local_norm.endswith(manifest_norm):
            best = max(best, 75)
        elif manifest_norm and manifest_norm in local_norm:
            best = max(best, 60)
    return best


def find_source_matches(source_dir: Path, groups: dict[tuple[str, str], DownloadGroup]) -> list[tuple[Path, DownloadGroup, int]]:
    candidates = [
        file_path
        for file_path in source_dir.iterdir()
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    matches: list[tuple[Path, DownloadGroup, int]] = []
    used_keys: set[tuple[str, str]] = set()
    for file_path in sorted(candidates, key=lambda path: path.stat().st_mtime, reverse=True):
        scored = [
            (score, key, group)
            for key, group in groups.items()
            if (score := match_score(file_path, group)) > 0
        ]
        if not scored:
            continue
        scored.sort(key=lambda item: item[0], reverse=True)
        score, key, group = scored[0]
        if key in used_keys:
            continue
        used_keys.add(key)
        matches.append((file_path, group, score))
    return matches


def extract_pdf_text(file_path: Path) -> tuple[str, str, int]:
    try:
        import pdfplumber  # type: ignore

        chunks = []
        with pdfplumber.open(str(file_path)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                chunks.append(f"## Page {index}\n\n{text}".strip())
            return "\n\n".join(chunks).strip(), "pdfplumber", len(pdf.pages)
    except Exception as first_error:
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(str(file_path))
            chunks = []
            for index, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                chunks.append(f"## Page {index}\n\n{text}".strip())
            return "\n\n".join(chunks).strip(), f"pypdf_after_pdfplumber_error:{type(first_error).__name__}", len(reader.pages)
        except Exception:
            raise first_error


def extract_docx_text(file_path: Path) -> tuple[str, str, int]:
    from docx import Document  # type: ignore

    document = Document(str(file_path))
    chunks = [paragraph.text for paragraph in document.paragraphs if clean_text(paragraph.text)]
    for table in document.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n\n".join(chunks).strip(), "python-docx", 0


def extract_text(file_path: Path) -> tuple[str, str, int]:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(file_path)
    if suffix == ".docx":
        return extract_docx_text(file_path)
    if suffix in {".txt", ".rtf"}:
        return file_path.read_text(encoding="utf-8", errors="replace"), "plain-text", 0
    raise ValueError(f"Unsupported file type: {suffix}")


def write_parsed_markdown(
    parsed_dir: Path,
    captured_file_id: str,
    group: DownloadGroup,
    repo_file_path: Path,
    text: str,
) -> Path:
    parsed_path = parsed_dir / f"{captured_file_id}.md"
    parsed_path.write_text(
        "\n".join(
            [
                f"# {repo_file_path.name}",
                "",
                f"- course_code: {group.course_code}",
                f"- course_id: {group.course_id}",
                f"- canvas_file_id: {group.canvas_file_id}",
                f"- repo_file_path: {repo_file_path}",
                f"- download_ids: {csv_join(group.download_ids)}",
                f"- download_manifest_ids: {csv_join(group.download_manifest_ids)}",
                f"- source_page_titles: {csv_join(group.source_page_titles)}",
                "",
                "## Extracted Text",
                "",
                text,
                "",
            ]
        ),
        encoding="utf-8",
    )
    return parsed_path


def ingest(args: argparse.Namespace) -> None:
    manifest_path = Path(args.manifest).expanduser().resolve()
    source_dir = Path(args.source_dir).expanduser().resolve()
    output_root = Path(args.output_root).expanduser().resolve()
    downloaded_dir = Path(args.downloaded_dir).expanduser().resolve()
    parsed_dir = Path(args.parsed_dir).expanduser().resolve()
    text_overrides = read_text_overrides(Path(args.text_overrides).expanduser().resolve())
    downloaded_dir.mkdir(parents=True, exist_ok=True)
    parsed_dir.mkdir(parents=True, exist_ok=True)

    manifest_rows = read_csv(manifest_path)
    groups = build_download_groups(manifest_rows)
    matches = find_source_matches(source_dir, groups)
    now = datetime.now(timezone.utc).isoformat()

    captured_rows: list[dict[str, object]] = []
    parsed_rows: list[dict[str, object]] = []
    parsed_by_manifest_id: dict[str, dict[str, object]] = {}

    for index, (source_path, group, score) in enumerate(matches, start=1):
        captured_file_id = f"captured-{index:04d}-{group.course_code.lower()}-{group.canvas_file_id}"
        destination_path = destination_path_for(source_path, downloaded_dir, group)
        if source_path.resolve() != destination_path.resolve():
            shutil.copy2(source_path, destination_path)
        digest = sha256_file(destination_path)
        parse_status = "parsed"
        parser = ""
        page_count = 0
        parsed_path = Path("")
        text = ""
        parse_error = ""
        try:
            text, parser, page_count = extract_text(destination_path)
            if len(substantive_text(text)) < 40:
                override = text_overrides.get((group.course_code, group.canvas_file_id))
                if override:
                    text = str(override.get("text") or "").strip()
                    parser = "manual-vision-override"
                    parse_status = clean_text(override.get("override_status")) or "parsed"
                    parse_error = clean_text(override.get("notes"))
                else:
                    parse_status = "parsed_text_empty_needs_ocr"
            parsed_path = write_parsed_markdown(parsed_dir, captured_file_id, group, destination_path, text)
        except Exception as error:
            override = text_overrides.get((group.course_code, group.canvas_file_id))
            if override:
                text = str(override.get("text") or "").strip()
                parser = "manual-vision-override"
                parse_status = clean_text(override.get("override_status")) or "parsed"
                parse_error = f"Applied override after {type(error).__name__}: {error}. {clean_text(override.get('notes'))}"
                parsed_path = write_parsed_markdown(parsed_dir, captured_file_id, group, destination_path, text)
            else:
                parse_status = "parse_failed"
                parse_error = f"{type(error).__name__}: {error}"

        captured_row = {
            "captured_file_id": captured_file_id,
            "course_id": group.course_id,
            "course_code": group.course_code,
            "canvas_file_id": group.canvas_file_id,
            "object_ids": csv_join(group.object_ids),
            "download_ids": csv_join(group.download_ids),
            "download_manifest_ids": csv_join(group.download_manifest_ids),
            "hrefs": csv_join(group.hrefs),
            "source_page_ids": csv_join(group.source_page_ids),
            "source_page_titles": csv_join(group.source_page_titles),
            "desktop_source_path": str(source_path),
            "repo_file_path": str(destination_path),
            "original_file_name": source_path.name,
            "manifest_file_names": csv_join(group.file_names),
            "file_type": destination_path.suffix.lower().lstrip("."),
            "byte_size": destination_path.stat().st_size,
            "sha256": digest,
            "match_score": score,
            "capture_status": "copied_to_repo",
            "parse_status": parse_status,
            "parsed_markdown_path": str(parsed_path) if parsed_path else "",
            "captured_at": now,
            "notes": parse_error,
        }
        captured_rows.append(captured_row)

        parsed_row = {
            "captured_file_id": captured_file_id,
            "course_id": group.course_id,
            "course_code": group.course_code,
            "canvas_file_id": group.canvas_file_id,
            "object_ids": csv_join(group.object_ids),
            "download_ids": csv_join(group.download_ids),
            "download_manifest_ids": csv_join(group.download_manifest_ids),
            "repo_file_path": str(destination_path),
            "parsed_markdown_path": str(parsed_path) if parsed_path else "",
            "parse_status": parse_status,
            "parser": parser,
            "file_type": destination_path.suffix.lower().lstrip("."),
            "page_count": page_count,
            "char_count": len(text),
            "sha256": digest,
            "title": source_path.name,
            "text_sample": clean_text(text)[:1000],
            "parsed_at": now,
            "notes": parse_error,
        }
        parsed_rows.append(parsed_row)
        for manifest_id in group.download_manifest_ids:
            parsed_by_manifest_id[manifest_id] = parsed_row

    execution_rows: list[dict[str, object]] = []
    for row in manifest_rows:
        parsed = parsed_by_manifest_id.get(clean_text(row.get("download_manifest_id")))
        if parsed:
            if parsed["parse_status"] == "parsed":
                execute_status = "downloaded_and_parsed"
                detail = "Matched a Desktop download, copied it into the repo, and parsed it."
            elif parsed["parse_status"] == "parsed_text_empty_needs_ocr":
                execute_status = "downloaded_needs_ocr"
                detail = "Matched a Desktop download, but text extraction was empty; needs OCR or vision review."
            else:
                execute_status = "downloaded_parse_failed"
                detail = parsed.get("notes", "")
        else:
            execute_status = "not_downloaded_browser_blocked"
            detail = "No matching downloaded file found in the Desktop landing folder yet."
        execution_rows.append({
            "download_manifest_id": row.get("download_manifest_id", ""),
            "download_id": row.get("download_id", ""),
            "course_id": row.get("course_id", ""),
            "course_code": row.get("course_code", ""),
            "object_id": row.get("object_id", ""),
            "file_name": row.get("file_name", ""),
            "canvas_file_id": row.get("canvas_file_id", ""),
            "href": row.get("href", ""),
            "execute_status": execute_status,
            "repo_file_path": parsed.get("repo_file_path", "") if parsed else "",
            "parsed_markdown_path": parsed.get("parsed_markdown_path", "") if parsed else "",
            "detail": detail,
        })

    write_csv(
        output_root / "downloaded_files_manifest.csv",
        captured_rows,
        [
            "captured_file_id",
            "course_id",
            "course_code",
            "canvas_file_id",
            "object_ids",
            "download_ids",
            "download_manifest_ids",
            "hrefs",
            "source_page_ids",
            "source_page_titles",
            "desktop_source_path",
            "repo_file_path",
            "original_file_name",
            "manifest_file_names",
            "file_type",
            "byte_size",
            "sha256",
            "match_score",
            "capture_status",
            "parse_status",
            "parsed_markdown_path",
            "captured_at",
            "notes",
        ],
    )
    write_csv(
        output_root / "parsed_downloads.csv",
        parsed_rows,
        [
            "captured_file_id",
            "course_id",
            "course_code",
            "canvas_file_id",
            "object_ids",
            "download_ids",
            "download_manifest_ids",
            "repo_file_path",
            "parsed_markdown_path",
            "parse_status",
            "parser",
            "file_type",
            "page_count",
            "char_count",
            "sha256",
            "title",
            "text_sample",
            "parsed_at",
            "notes",
        ],
    )
    write_csv(
        output_root / "download_execution_status.csv",
        execution_rows,
        [
            "download_manifest_id",
            "download_id",
            "course_id",
            "course_code",
            "object_id",
            "file_name",
            "canvas_file_id",
            "href",
            "execute_status",
            "repo_file_path",
            "parsed_markdown_path",
            "detail",
        ],
    )

    print(f"Matched and copied {len(captured_rows)} file(s) from {source_dir} into {downloaded_dir}")
    parsed_count = sum(1 for row in parsed_rows if row["parse_status"] == "parsed")
    print(f"Parsed {parsed_count} file(s); {len(manifest_rows) - len(parsed_by_manifest_id)} manifest row(s) remain without repo file content")


def parse_args() -> argparse.Namespace:
    repo_root = Path(__file__).resolve().parents[2]
    output_root = repo_root / "canvas" / "indexer" / "output"
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--manifest", default=str(output_root / "download_manifest.csv"))
    parser.add_argument("--source-dir", default=str(Path.home() / "Desktop"))
    parser.add_argument("--output-root", default=str(output_root))
    parser.add_argument("--downloaded-dir", default=str(output_root / "downloaded_files"))
    parser.add_argument("--parsed-dir", default=str(output_root / "parsed_files"))
    parser.add_argument("--text-overrides", default=str(repo_root / "canvas" / "indexer" / "input" / "download_text_overrides.csv"))
    return parser.parse_args()


if __name__ == "__main__":
    ingest(parse_args())
